print(">>> fill_template.py התחיל לרוץ")

import os
import json
import re
from docx import Document

from google import genai
from google.genai import types


# =========================
# 1) Gemini key (ENV בלבד)
# =========================
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
if not GEMINI_API_KEY:
    raise RuntimeError('חסר GEMINI_API_KEY ב-Environment. (PowerShell: setx GEMINI_API_KEY "YOUR_KEY")')

client = genai.Client(api_key=GEMINI_API_KEY)


# =========================
# 2) Helpers
# =========================
def rtl(text: str) -> str:
    """עוטף טקסט בכיווניות RTL כדי שסימני פיסוק/מספרים לא יברחו."""
    if not text:
        return text
    return "\u202B" + str(text) + "\u202C"  # RLE ... PDF


def strip_prices(line: str) -> str:
    """
    מוריד מחירים/מספרים מסוף סעיף.
    שימושי אם רוצים להציג סעיפים בלי מחירים (רק סה"כ).
    """
    if not line:
        return ""
    line = str(line).strip()

    # מוריד "- 4000", "– 4,000 ₪" בסוף שורה
    line = re.sub(r'[\-\–]\s*[\d,.\s]+₪?\s*$', '', line).strip()
    # מוריד "(4000)" בסוף שורה
    line = re.sub(r'\(\s*[\d,.\s]+₪?\s*\)\s*$', '', line).strip()
    # אם נשאר מספר גדול מנותק, מוחק אותו (שמרני)
    line = re.sub(r'\b[\d,]{3,}\b', '', line).strip()

    line = re.sub(r'\s{2,}', ' ', line).strip()
    return line


def replace_placeholders_in_paragraph(paragraph, values: dict):
    """
    החלפה ברמת runs כדי לשמור עיצוב.
    הערה: יכול להיכשל אם placeholder מפוצל ליותר מ-run אחד בתבנית.
    """
    for placeholder, value in values.items():
        if placeholder in paragraph.text:
            for run in paragraph.runs:
                if placeholder in run.text:
                    run.text = run.text.replace(placeholder, value)


def replace_placeholders_everywhere(doc: Document, values: dict):
    # paragraphs
    for p in doc.paragraphs:
        replace_placeholders_in_paragraph(p, values)

    # tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_placeholders_in_paragraph(p, values)


# =========================
# 3) Gemini: ניסוח תיאור + תנאי תשלום בלבד
# =========================
def process_quote_with_ai(raw_data: dict) -> dict:
    system_msg = (
        "אתה כותב הצעות מחיר מקצועיות בעברית עבור קבלן שיפוצים. "
        "תפקידך הוא רק לסדר ניסוח יפה וברור של תיאור העבודה ותנאי התשלום. "
        "אסור לך להמציא מחירים, לשנות סכומים, להוסיף סעיפים או למחוק סעיפים. "
        "תחזיר אך ורק JSON תקין, ללא טקסט נוסף."
    )

    user_msg = f"""
שם הלקוח: {raw_data.get('client_name', '')}
כתובת: {raw_data.get('address', '')}
סוג עבודה: {raw_data.get('job_type', '')}

תיאור גולמי של העבודה: {raw_data.get('raw_description', '')}

תנאי תשלום (גולמי): {raw_data.get('payment_terms', '')}

תחזיר JSON בלבד במבנה:
{{
  "work_description": "תיאור כללי קצר ומקצועי של סוג העבודה.",
  "payment_terms": "ניסוח מסודר ומקצועי של תנאי התשלום, בלי לשנות את המשמעות."
}}
"""

    schema = {
        "type": "OBJECT",
        "required": ["work_description", "payment_terms"],
        "properties": {
            "work_description": {"type": "STRING"},
            "payment_terms": {"type": "STRING"},
        },
    }

    resp = client.models.generate_content(
        model="gemini-2.5-flash",
        contents=user_msg,
        config=types.GenerateContentConfig(
            system_instruction=system_msg,
            response_mime_type="application/json",
            response_schema=schema,
            temperature=0.2,
        ),
    )

    if getattr(resp, "parsed", None) is not None:
        return resp.parsed

    # fallback אם חזר טקסט (נדיר)
    raw_text = (resp.text or "").strip()
    raw_text = re.sub(r"^```[a-zA-Z0-9]*\s*", "", raw_text)
    raw_text = re.sub(r"\s*```$", "", raw_text)
    start = raw_text.find("{")
    end = raw_text.rfind("}")
    cleaned = raw_text[start:end + 1] if start != -1 and end != -1 else raw_text
    return json.loads(cleaned)


# =========================
# 4) בניית סעיפים
# =========================
def build_price_section(job_type: str, raw_price_lines: list, show_line_prices: bool = True) -> str:
    job_type = (job_type or "").strip()
    title = f"עבודות {job_type}:" if job_type else "עבודות:"

    lines = [title]
    RLM = "\u200F"  # עוזר עם נקודות/מקפים בעברית

    for line in raw_price_lines or []:
        s = str(line).strip()
        if not s:
            continue

        clean = s if show_line_prices else strip_prices(s)
        if clean:
            lines.append(f"{RLM}• {clean}")

    return "\n".join(lines)


# =========================
# 5) main: מילוי תבנית
# =========================
def fill_template(template_path: str, output_path: str, raw_data: dict):
    processed = process_quote_with_ai(raw_data)

    show_line_prices = raw_data.get("show_line_prices", True)  # ברירת מחדל: כן להציג מחירים בסעיפים
    price_section = build_price_section(
        raw_data.get("job_type", ""),
        raw_data.get("raw_price_lines", []),
        show_line_prices=show_line_prices,
    )

    total_price_str = str(raw_data.get("total_price", "") or "").replace("₪", "").replace(",", "").strip()

    values = {
        "{{CLIENT_NAME}}": rtl(raw_data.get("client_name", "")),
        "{{ADDRESS}}": rtl(raw_data.get("address", "")),
        "{{PRICE_SECTION}}": rtl(price_section),
        "{{TOTAL_PRICE}}": rtl(total_price_str),
        "{{PAYMENT_TERMS}}": rtl(processed.get("payment_terms", "")),
        "{{WORK_DESCRIPTION}}": rtl(processed.get("work_description", "")),  # אם לא קיים בתבנית, פשוט לא יוחלף
    }

    doc = Document(template_path)
    replace_placeholders_everywhere(doc, values)
    doc.save(output_path)


# =========================
# בדיקה מקומית
# =========================
if __name__ == "__main__":
    raw_data = {
        "client_name": "משפחת כהן",
        "address": "החרצית 20 חולון, דירה 22 קומה 7",
        "job_type": "שיפוץ דירה – מטבח וריצוף",
        "raw_description": "שיפוץ מטבח וחללי הבית כולל פירוק, העתקת נקודות מים וחשמל, ריצוף חדש וצביעה.",
        "raw_price_lines": [
            "פירוק מטבח ישן כולל פינוי – 2500 ₪",
            "העתקת נקודות מים + ניקוז – 1500 ₪",
            "צביעה כללית לבית – 12000 ₪",
            "פינוי פסולת ע\"י הקבלן",
        ],
        "payment_terms": "התשלום יתבצע לפי התקדמות העבודה. המחיר אינו כולל מע\"מ.",
        "total_price": "49500",
        "show_line_prices": True,  # שנה ל-False אם אתה רוצה להסתיר מחירי סעיפים
    }

    fill_template("template.docx", "הצעת_מחיר_בדיקה.docx", raw_data)
    print(">>> נוצר: הצעת_מחיר_בדיקה.docx")
